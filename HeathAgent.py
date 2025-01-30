import streamlit as st
from crewai import Agent, Task, Crew, Process
from langchain_community.llms import Ollama
from docx import Document
from io import BytesIO
import base64
import os
import time

# Configure Ollama
os.environ["OLLAMA_API_BASE"] = "http://localhost:11435"

# Custom CSS styling
st.markdown("""
<style>
    .header {
        padding: 1rem;
        background: linear-gradient(135deg, #6366f1, #8b5cf6);
        color: white;
        border-radius: 10px;
        margin-bottom: 2rem;
    }
    .input-card {
        padding: 1.5rem;
        background: #f8f9fa;
        border-radius: 10px;
        box-shadow: 0 2px 4px rgba(0,0,0,0.1);
        margin-bottom: 1.5rem;
    }
    .chat-card {
        padding: 1rem;
        background: white;
        border-radius: 10px;
        box-shadow: 0 2px 4px rgba(0,0,0,0.05);
        margin: 0.5rem 0;
        max-height: 500px;
        overflow-y: auto;
    }
    .user-message {
        background: #eef2ff;
        padding: 0.8rem;
        border-radius: 10px;
        margin: 0.5rem 0;
    }
    .bot-message {
        background: #f0fdf4;
        padding: 0.8rem;
        border-radius: 10px;
        margin: 0.5rem 0;
    }
</style>
""", unsafe_allow_html=True)

def generate_docx(result):
    doc = Document()
    doc.add_heading('HealthAgent Report', 0)
    doc.add_paragraph(result)
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

def get_download_link(bio, filename):
    b64 = base64.b64encode(bio.read()).decode()
    return f'<a class="download-btn" href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="{filename}">📥 Download Full Report</a>'

# Initialize Ollama
llm = Ollama(
    model="qwen2.5:1.5b",
    temperature=0.1,
    timeout=600
)

# Initialize session state
if 'chat_history' not in st.session_state:
    st.session_state.chat_history = []

# App Header
st.markdown("""
<div class="header">
    <h1 style="color:white; margin:0;">🩺 HealthAgent</h1>
    <p style="color:white; margin:0; opacity:0.9;">AI-Powered Medical Analysis & Health Companion</p>
</div>
""", unsafe_allow_html=True)

# Create tabs
tab1, tab2 = st.tabs(["Medical Analysis", "Health Chatbot"])

with tab1:
    # Medical Analysis Section
    with st.container():
        with st.form("patient_info"):
            st.markdown('<div class="input-card">', unsafe_allow_html=True)
            
            col1, col2 = st.columns(2)
            with col1:
                gender = st.selectbox("Gender", ("Male", "Female", "Other"), index=0)
                age = st.number_input("Age", min_value=0, max_value=120, value=25)
            
            with col2:
                symptoms = st.text_area("Current Symptoms", placeholder="Headache, fever, cough...", height=100)
                medical_history = st.text_area("Medical History", placeholder="Previous conditions, allergies...", height=100)
            
            st.markdown('</div>', unsafe_allow_html=True)
            
            submitted = st.form_submit_button("🚀 Analyze My Health", use_container_width=True)

    # Agent Definitions for Medical Analysis
    diagnostician = Agent(
        role="Senior Diagnostician",
        goal="Identify potential medical conditions",
        backstory="Expert physician with 20+ years experience",
        verbose=True,
        allow_delegation=False,
        llm=llm
    )

    treatment_advisor = Agent(
        role="Medical Advisor",
        goal="Provide health recommendations",
        backstory="Combines evidence-based medicine with holistic approaches",
        verbose=True,
        llm=llm
    )

    # Task Definitions
    diagnosis_task = Task(
        description=f"""
        Analyze patient profile:
        - Gender: {gender}
        - Age: {age}
        - Symptoms: {symptoms}
        - Medical History: {medical_history}
        Output format:
        1. List 3 possible conditions with probability percentages
        2. Flag urgent risks in bold
        3. Explain reasoning simply""",
        agent=diagnostician,
        expected_output="Structured report with diagnoses, probabilities, and explanations"
    )

    advice_task = Task(
        description=f"""
        Create recommendations including:
        1. Immediate actions
        2. Lifestyle changes
        3. When to see a doctor
        4. Preventive measures""",
        agent=treatment_advisor,
        expected_output="Numbered list of actionable recommendations"
    )

    medical_crew = Crew(
        agents=[diagnostician, treatment_advisor],
        tasks=[diagnosis_task, advice_task],
        verbose=1,
        process=Process.sequential
    )

    # Process Analysis
    if submitted:
        if not symptoms.strip():
            st.warning("Please describe your symptoms")
        else:
            with st.status("🔍 Analyzing your health profile...", expanded=True) as status:
                try:
                    start_time = time.time()
                    result = medical_crew.kickoff()
                    
                    st.markdown('<div class="result-card">', unsafe_allow_html=True)
                    st.subheader("Health Analysis")
                    st.success("Analysis completed successfully!")
                    st.markdown("---")
                    st.write(result)
                    st.markdown("---")
                    
                    with st.spinner("Generating report..."):
                        docx_file = generate_docx(result)
                        st.markdown(
                            get_download_link(docx_file, "healthagent_report.docx"),
                            unsafe_allow_html=True
                        )
                    
                    st.markdown('</div>', unsafe_allow_html=True)
                    st.balloons()
                    
                except Exception as e:
                    st.error(f"Error: {str(e)}")
                    st.info("Please check: 1. Ollama is running 2. Model is downloaded")

with tab2:
    # Health Chatbot Section
    st.markdown("### 💬 Health Companion Chat")
    st.caption("Ask about nutrition, exercise, medical terms, or general health advice")
    
    # Chat container
    chat_container = st.container()
    
    # User input
    user_query = st.chat_input("Type your health question here...")
    
    if user_query:
        # Add user message to history
        st.session_state.chat_history.append({"role": "user", "content": user_query})
        
        # Generate bot response
        with st.spinner("HealthAgent is thinking..."):
            try:
                response = llm.invoke(f"""
                [SYSTEM PROMPT]
                You are a helpful health assistant. Provide accurate, practical advice about:
                - Nutrition and diet plans
                - Exercise routines
                - Medical terminology
                - General wellness tips
                - Healthy lifestyle recommendations
                
                Keep responses concise and easy to understand. Use bullet points when appropriate.
                If unsure, recommend consulting a healthcare professional.
                
                [USER QUESTION]
                {user_query}
                """)
                
                st.session_state.chat_history.append({"role": "bot", "content": response})
                
            except Exception as e:
                st.error(f"Error generating response: {str(e)}")
    
    # Display chat history
    with chat_container:
        for message in st.session_state.chat_history[-5:]:  # Show last 5 messages
            if message["role"] == "user":
                st.markdown(f'<div class="user-message">👤 You: {message["content"]}</div>', 
                           unsafe_allow_html=True)
            else:
                st.markdown(f'<div class="bot-message">🤖 HealthAgent: {message["content"]}</div>', 
                           unsafe_allow_html=True)

# Footer
st.markdown("---")
st.caption("""
⚠️ **Important Notice**: HealthAgent provides preliminary health insights only. 
Always consult a qualified healthcare professional for medical decisions.
""")

# Sidebar
with st.sidebar:
    st.markdown("## About HealthAgent")
    st.markdown("""
    Your comprehensive health AI that:
    - Analyzes medical symptoms
    - Provides health recommendations
    - Answers general health questions
    - Offers fitness/nutrition advice
    """)
    st.markdown("---")
    st.markdown("**How to use:**")
    st.markdown("""
    1. Medical Analysis tab for symptom checking
    2. Chatbot tab for general health queries
    3. Download reports anytime
    """)
    st.markdown("---")
    st.markdown("🛠️ Powered by:")
    st.markdown("- Ollama LLM")
    st.markdown("- Qwen2.5 1.5B model")
    st.markdown("- Streamlit UI")